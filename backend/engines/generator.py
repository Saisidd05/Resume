"""
Resume Generation Engine — fills the template with user answers.

CARDINAL RULE: The output file must be visually indistinguishable from the
input template, except for replaced content.

Strategy by file type:
  PDF:
    - Use PyMuPDF to open the original PDF.
    - For each placeholder bounding box, cover it with a white rectangle
      (matching the page background), then draw the new text at the exact
      same position using the same font size and alignment.
    - Prefer redacting and re-drawing over simple overlay to avoid artifacts.

  DOCX:
    - Open the original DOCX with python-docx.
    - Walk through paragraphs and runs, replacing placeholder text while
      preserving ALL run formatting (font, size, bold, italic, color, spacing).
    - Never touch paragraph or table structure.
"""
from __future__ import annotations

import base64
import logging
import re
from io import BytesIO
from typing import Optional

import fitz  # PyMuPDF

from backend.models.schemas import (
    GenerationRequest, GenerationResult, FileType,
    Placeholder, AnswerSet, FieldAnswer, ParsedTemplate
)
from backend.engines.template_engine import (
    get_section_constraints, enforce_constraints
)

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
# Answer lookup helper
# ─────────────────────────────────────────────────────────────────────────────

def _build_answer_map(answers: AnswerSet) -> dict[str, str]:
    """Build a {field_id: value} lookup from the answer set."""
    return {a.field_id: a.value for a in answers.answers}


def _get_answer_for_placeholder(
    ph: Placeholder,
    answer_map: dict[str, str],
    section_answers: dict[str, list[FieldAnswer]],
) -> Optional[str]:
    """
    Find the best matching answer for a placeholder.
    Tries: exact field_id link → label keyword match → section aggregate.
    """
    # 1. Check if any answer field is directly linked to this placeholder
    for field_id, value in answer_map.items():
        # Field IDs are like "section_0_full_name" — check placeholder label
        ph_label_lower = ph.label.lower().replace(' ', '_')
        if ph_label_lower in field_id.lower():
            return value

    # 2. Try keyword matching between placeholder label and answer field IDs
    ph_words = set(ph.label.lower().split())
    best_match: Optional[str] = None
    best_score = 0

    for field_id, value in answer_map.items():
        field_words = set(field_id.lower().replace('_', ' ').split())
        overlap = len(ph_words & field_words)
        if overlap > best_score:
            best_score = overlap
            best_match = value

    if best_score >= 1 and best_match:
        return best_match

    # 3. Return first non-empty answer from the same section
    sec_answers = section_answers.get(ph.section, [])
    for ans in sec_answers:
        if ans.value.strip():
            return ans.value

    return None


def _build_section_answers(answers: AnswerSet) -> dict[str, list[FieldAnswer]]:
    """Group answers by section_id."""
    result: dict[str, list[FieldAnswer]] = {}
    for ans in answers.answers:
        result.setdefault(ans.section_id, []).append(ans)
    return result


# ─────────────────────────────────────────────────────────────────────────────
# PDF Generator — text overlay approach
# ─────────────────────────────────────────────────────────────────────────────

def _generate_pdf(request: GenerationRequest) -> GenerationResult:
    """
    Fill a PDF template by:
    1. Opening the original PDF from base64.
    2. For each placeholder, covering the old text with a white box and
       drawing the new answer text at the exact same bounding box.
    3. Matching font size from the placeholder's recorded FontInfo.
    """
    errors: list[str] = []
    warnings: list[str] = []

    try:
        pdf_bytes = base64.b64decode(request.original_file_b64)
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as e:
        return GenerationResult(
            success=False,
            errors=[f"Failed to open template PDF: {e}"]
        )

    answer_map = _build_answer_map(request.answers)
    section_answers = _build_section_answers(request.answers)

    # Group placeholders by page for efficient processing
    placeholders_by_page: dict[int, list[Placeholder]] = {}
    for ph in request.parsed_template.all_placeholders:
        if ph.bbox is not None:
            page_num = ph.bbox.page
            placeholders_by_page.setdefault(page_num, []).append(ph)

    for page_num, page_phs in placeholders_by_page.items():
        if page_num >= len(doc):
            warnings.append(f"Page {page_num} not found in document")
            continue

        page = doc[page_num]

        for ph in page_phs:
            new_text = _get_answer_for_placeholder(ph, answer_map, section_answers)

            if not new_text:
                warnings.append(f"No answer found for placeholder '{ph.label}', keeping original")
                continue

            # Apply constraints if strict mode
            if request.strict_mode:
                section = next(
                    (s for s in request.parsed_template.sections if s.id == ph.section),
                    None
                )
                if section:
                    constraints = get_section_constraints(section, request.parsed_template)
                    new_text = enforce_constraints(new_text, constraints)

            # Get bounding box
            bbox = ph.bbox
            rect = fitz.Rect(bbox.x0, bbox.y0, bbox.x1, bbox.y1)

            # Expand rect slightly for comfortable text fit
            text_rect = fitz.Rect(
                bbox.x0,
                bbox.y0 - 2,
                bbox.x1 + 50,  # allow some horizontal overflow
                bbox.y1 + 4,
            )

            # ── Cover old placeholder text with white rectangle ────────────
            # This preserves lines/borders around the placeholder
            page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))

            # ── Determine font size ────────────────────────────────────────
            font_size = ph.font.size if ph.font.size > 0 else 11.0

            # ── Choose font ────────────────────────────────────────────────
            # PyMuPDF built-in fonts: helv, tiro, cour, zadb, symb
            font_name = "helv"  # Helvetica — closest universal match
            if ph.font.bold and ph.font.italic:
                font_name = "hebo"  # Helvetica Bold Oblique
            elif ph.font.bold:
                font_name = "hebo"  # Helvetica Bold
            elif ph.font.italic:
                font_name = "heit"  # Helvetica Italic

            # ── Parse color ────────────────────────────────────────────────
            color = (0, 0, 0)  # Default black
            if ph.font.color and ph.font.color != "#000000":
                try:
                    hex_c = ph.font.color.lstrip('#')
                    r, g, b = (int(hex_c[i:i+2], 16) / 255 for i in (0, 2, 4))
                    color = (r, g, b)
                except Exception:
                    pass

            # ── Draw new text ──────────────────────────────────────────────
            try:
                # Handle multi-line text for paragraph placeholders
                lines = new_text.split('\n')
                y_pos = bbox.y0 + font_size  # baseline position

                for line in lines:
                    if line.strip():
                        page.insert_text(
                            point=fitz.Point(bbox.x0, y_pos),
                            text=line.strip(),
                            fontname=font_name,
                            fontsize=font_size,
                            color=color,
                        )
                    y_pos += font_size * 1.4  # line height factor

            except Exception as e:
                errors.append(f"Failed to write text for '{ph.label}': {e}")
                logger.warning(f"Text insertion failed for placeholder {ph.id}: {e}")

    # ── Save to bytes ──────────────────────────────────────────────────────────
    try:
        output_stream = BytesIO()
        doc.save(output_stream)
        doc.close()
        output_bytes = output_stream.getvalue()
        output_b64 = base64.b64encode(output_bytes).decode("utf-8")

        output_filename = request.original_filename.replace(
            '.pdf', '_filled.pdf'
        ).replace('.PDF', '_filled.pdf')

        return GenerationResult(
            success=True,
            output_file_b64=output_b64,
            output_filename=output_filename,
            file_type=FileType.PDF,
            errors=errors,
            warnings=warnings,
        )
    except Exception as e:
        return GenerationResult(
            success=False,
            errors=[f"Failed to save output PDF: {e}"] + errors,
        )


# ─────────────────────────────────────────────────────────────────────────────
# DOCX Generator — run-level replacement
# ─────────────────────────────────────────────────────────────────────────────

def _generate_docx(request: GenerationRequest) -> GenerationResult:
    """
    Fill a DOCX template by walking runs and replacing placeholder text
    while preserving every formatting attribute on the run.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    errors: list[str] = []
    warnings: list[str] = []

    try:
        docx_bytes = base64.b64decode(request.original_file_b64)
        doc = Document(BytesIO(docx_bytes))
    except Exception as e:
        return GenerationResult(
            success=False,
            errors=[f"Failed to open template DOCX: {e}"]
        )

    answer_map = _build_answer_map(request.answers)
    section_answers = _build_section_answers(request.answers)

    # Build a lookup: paragraph_index → list of placeholders
    ph_by_para: dict[int, list[Placeholder]] = {}
    for ph in request.parsed_template.all_placeholders:
        if ph.paragraph_index is not None:
            ph_by_para.setdefault(ph.paragraph_index, []).append(ph)

    def replace_in_paragraph(para, para_idx: int) -> None:
        """Replace placeholder text in a paragraph's runs."""
        phs = ph_by_para.get(para_idx, [])
        if not phs:
            return

        for ph in phs:
            new_text = _get_answer_for_placeholder(ph, answer_map, section_answers)
            if not new_text:
                continue

            # Apply constraints
            if request.strict_mode:
                section = next(
                    (s for s in request.parsed_template.sections if s.id == ph.section),
                    None
                )
                if section:
                    constraints = get_section_constraints(section, request.parsed_template)
                    new_text = enforce_constraints(new_text, constraints)

            # Replace in the specific run if run_index is known
            if ph.run_index is not None and ph.run_index < len(para.runs):
                run = para.runs[ph.run_index]
                # Preserve formatting — only replace text
                run.text = new_text
            else:
                # Full-paragraph replacement: find the run containing the placeholder text
                full_text = para.text
                if ph.original_text in full_text:
                    # Replace across potentially multiple runs
                    _replace_across_runs(para, ph.original_text, new_text)

    def _replace_across_runs(para, old_text: str, new_text: str) -> None:
        """
        Replace old_text with new_text across runs in a paragraph.
        Consolidates text into the first matching run to preserve formatting.
        """
        # Build concatenated text with run boundaries
        run_texts = [r.text for r in para.runs]
        full = "".join(run_texts)

        if old_text not in full:
            return

        new_full = full.replace(old_text, new_text, 1)

        # Redistribute new text across runs preserving original run lengths
        pos = 0
        for run in para.runs:
            run_len = len(run_texts[para.runs.index(run)])
            run.text = new_full[pos:pos + run_len]
            pos += run_len

        # If new text is longer, put remainder in last run
        if pos < len(new_full):
            para.runs[-1].text += new_full[pos:]

    # Walk all paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        replace_in_paragraph(para, para_idx)

    # Walk table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, para_idx=999999)

    # Save to bytes
    try:
        output_stream = BytesIO()
        doc.save(output_stream)
        output_bytes = output_stream.getvalue()
        output_b64 = base64.b64encode(output_bytes).decode("utf-8")

        output_filename = request.original_filename.replace(
            '.docx', '_filled.docx'
        ).replace('.DOCX', '_filled.docx')

        return GenerationResult(
            success=True,
            output_file_b64=output_b64,
            output_filename=output_filename,
            file_type=FileType.DOCX,
            errors=errors,
            warnings=warnings,
        )
    except Exception as e:
        return GenerationResult(
            success=False,
            errors=[f"Failed to save output DOCX: {e}"] + errors,
        )


# ─────────────────────────────────────────────────────────────────────────────
# Main entry point
# ─────────────────────────────────────────────────────────────────────────────

def generate_resume(request: GenerationRequest) -> GenerationResult:
    """Dispatch to the correct generator based on file type."""
    if request.file_type == FileType.PDF:
        return _generate_pdf(request)
    elif request.file_type == FileType.DOCX:
        return _generate_docx(request)
    else:
        return GenerationResult(
            success=False,
            errors=[f"Unsupported file type: {request.file_type}"]
        )
