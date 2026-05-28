"""
DOCX Parser — extracts structure from resume template Word documents.

Strategy:
  1. Use python-docx to traverse paragraphs and table cells.
  2. Identify section headings via paragraph styles (Heading1/2/3, bold large fonts).
  3. Extract placeholder text using the same regex patterns as PDF parser.
  4. Record paragraph + run indices so the generator can precisely replace content
     while preserving ALL formatting (bold, italic, font size, color, spacing).
  5. Extract written instructions from paragraph text.

CRITICAL: Parser reads only — never writes to the original document.
"""
from __future__ import annotations

import re
import base64
import logging
from typing import Optional
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from backend.models.schemas import (
    ParsedTemplate, Section, Placeholder, FontInfo,
    TemplateInstruction, FileType
)
from backend.parsers.pdf_parser import (
    PLACEHOLDER_PATTERNS, _extract_instruction, _infer_label, _is_placeholder_text
)

logger = logging.getLogger(__name__)

# Styles that indicate a heading paragraph
HEADING_STYLE_NAMES = {
    'heading 1', 'heading 2', 'heading 3',
    'heading1', 'heading2', 'heading3',
    'title', 'subtitle',
}


def _run_font_info(run) -> FontInfo:
    """Extract FontInfo from a python-docx Run object."""
    try:
        font = run.font
        name = font.name or "Calibri"
        size = font.size.pt if font.size else 11.0
        bold = bool(font.bold)
        italic = bool(font.italic)
        color_str = "#000000"
        if font.color and font.color.type is not None:
            try:
                rgb = font.color.rgb
                color_str = f"#{rgb}"
            except Exception:
                pass
        return FontInfo(name=name, size=round(size, 2), bold=bold, italic=italic, color=color_str)
    except Exception:
        return FontInfo()


def _paragraph_font_info(para) -> FontInfo:
    """Best-effort FontInfo from paragraph's first run (or paragraph style)."""
    for run in para.runs:
        if run.text.strip():
            return _run_font_info(run)
    # Fallback: use paragraph style font
    try:
        style = para.style
        if style and style.font:
            return FontInfo(
                name=style.font.name or "Calibri",
                size=style.font.size.pt if style.font.size else 11.0,
                bold=bool(style.font.bold),
                italic=bool(style.font.italic),
            )
    except Exception:
        pass
    return FontInfo()


def _is_heading_paragraph(para) -> bool:
    """Determine if a paragraph is a section heading."""
    style_name = (para.style.name or "").lower()

    # Explicit heading style
    if any(h in style_name for h in HEADING_STYLE_NAMES):
        return True

    # Bold, larger font, short text (typical for section headers)
    text = para.text.strip()
    if not text or len(text) > 80:
        return False

    fi = _paragraph_font_info(para)
    if fi.bold and fi.size >= 12:
        # Also check it's not just a name (usually the biggest font)
        if len(text.split()) <= 6 and text == text.upper():
            return True  # ALL-CAPS bold short text = heading
        if len(text.split()) <= 4 and fi.size >= 14:
            return True

    return False


def _is_bullet_paragraph(para) -> bool:
    """Check if paragraph is a bullet/list item."""
    style_name = (para.style.name or "").lower()
    if 'list' in style_name or 'bullet' in style_name:
        return True
    # Check for bullet characters
    text = para.text.strip()
    if text and text[0] in ('•', '◦', '▪', '-', '→', '*'):
        return True
    return False


def parse_docx(file_bytes: bytes, filename: str) -> ParsedTemplate:
    """
    Main DOCX parsing entry point.
    Returns a fully populated ParsedTemplate.
    """
    doc = Document(BytesIO(file_bytes))
    original_b64 = base64.b64encode(file_bytes).decode("utf-8")

    sections: list[Section] = []
    current_section: Optional[Section] = None
    all_instructions: list[TemplateInstruction] = []
    all_placeholders: list[Placeholder] = []
    placeholder_counter = 0

    def process_paragraph(para, para_idx: int) -> None:
        """Process a single paragraph — detect heading, instructions, placeholders."""
        nonlocal current_section, placeholder_counter

        text = para.text.strip()
        if not text:
            if current_section:
                current_section.raw_content += "\n"
            return

        fi = _paragraph_font_info(para)

        # ── Check for section heading ─────────────────────────────────────
        if _is_heading_paragraph(para):
            current_section = Section(
                id=f"section_{len(sections)}",
                name=text,
                order=len(sections),
                heading_text=text,
                heading_font=fi,
            )
            sections.append(current_section)
            return

        # Ensure we have a section
        if current_section is None:
            current_section = Section(
                id="section_header",
                name="Personal Information",
                order=0,
                heading_text="Personal Information",
                heading_font=fi,
            )
            sections.insert(0, current_section)

        current_section.raw_content += text + "\n"

        # ── Check for written instructions ────────────────────────────────
        instruction = _extract_instruction(text, current_section.name)
        if instruction:
            current_section.instructions.append(instruction)
            all_instructions.append(instruction)

        # ── Check each run for placeholder content ────────────────────────
        is_bullet = _is_bullet_paragraph(para)
        for run_idx, run in enumerate(para.runs):
            run_text = run.text
            if _is_placeholder_text(run_text):
                run_fi = _run_font_info(run)
                ph = Placeholder(
                    id=f"ph_{placeholder_counter}",
                    label=_infer_label(run_text.strip(), current_section.name),
                    section=current_section.id,
                    original_text=run_text.strip(),
                    font=run_fi,
                    paragraph_index=para_idx,
                    run_index=run_idx,
                    is_bullet=is_bullet,
                    is_heading=False,
                )
                current_section.placeholders.append(ph)
                all_placeholders.append(ph)
                placeholder_counter += 1

        # ── Also check full paragraph text as a placeholder ───────────────
        if _is_placeholder_text(text) and not any(
            ph.paragraph_index == para_idx for ph in all_placeholders
        ):
            ph = Placeholder(
                id=f"ph_{placeholder_counter}",
                label=_infer_label(text, current_section.name),
                section=current_section.id,
                original_text=text,
                font=fi,
                paragraph_index=para_idx,
                run_index=None,
                is_bullet=is_bullet,
                is_heading=False,
            )
            current_section.placeholders.append(ph)
            all_placeholders.append(ph)
            placeholder_counter += 1

    # ── Walk paragraphs ───────────────────────────────────────────────────
    for para_idx, para in enumerate(doc.paragraphs):
        process_paragraph(para, para_idx)

    # ── Walk table cells ─────────────────────────────────────────────────
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    # Use a large index offset to avoid collision with main paragraphs
                    process_paragraph(para, para_idx=999999)

    # ── Re-index ─────────────────────────────────────────────────────────
    for i, sec in enumerate(sections):
        sec.order = i

    # Default font: most common style in the document
    default_font = FontInfo(name="Calibri", size=11.0)
    try:
        style = doc.styles['Normal']
        if style.font.name:
            default_font.name = style.font.name
        if style.font.size:
            default_font.size = style.font.size.pt
    except Exception:
        pass

    return ParsedTemplate(
        file_type=FileType.DOCX,
        original_filename=filename,
        page_count=1,  # DOCX page count requires rendering
        sections=sections,
        all_instructions=all_instructions,
        all_placeholders=all_placeholders,
        original_file_b64=original_b64,
        default_font=default_font,
    )
